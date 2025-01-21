from docx import Document
import os

document = Document ()

document.add_heading('Titulo del trabajo', level=1)

document.add_paragraph('Esta es la introduccion del trabajo')

document.add_heading('Desarrollo', level=2)

document.add_paragraph('Aqui se desarrolla el contenido principal del trabajo')

document.add_heading('Conclusion', level=2)

document.add_paragraph('Esta es la conclusion del trabajo')

document.save('Trabajo.docx')

print("Se ha creado el archivo 'Trabajo.docx'")

os.system("start Trabajo.docx")