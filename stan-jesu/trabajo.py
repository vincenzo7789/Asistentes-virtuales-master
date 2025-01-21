# This python file uses the following encoding: utf-8 
from docx import Document
from docx.shared import Inches
import os

def crear_trabajo_fotosintesis():
    document = Document()

    # Agregar titulo
    document.add_heading('Fotosintesis', 0)

    # Agregar parrafo introductorio
    document.add_paragraph('La fotosintesis es el proceso mediante el cual las plantas verdes y otros organismos utilizan la energia de la luz solar para convertir el dioxido de carbono y el agua en glucosa (un azucar) y oxigeno.')

    # Agregar una seccion sobre las etapas de la fotosintesis
    document.add_heading('Etapas de la Fotosintesis', level=1)
    document.add_paragraph('La fotosintesis se divide en dos etapas principales:')

    # Etapa 1: Fase luminosa
    document.add_paragraph('Fase luminosa: Esta etapa ocurre en las membranas de los tilacoides de los cloroplastos. La energia luminosa es absorbida por la clorofila y otros pigmentos, generando ATP y NADPH.  Se libera oxigeno como subproducto.')

    # Etapa 2: Fase oscura
    document.add_paragraph('Fase oscura (ciclo de Calvin): Esta etapa ocurre en el estroma de los cloroplastos.  El ATP y el NADPH producidos en la fase luminosa se utilizan para convertir el dioxido de carbono en glucosa mediante una serie de reacciones enzimaticas.')


    # Agregar una seccion sobre la importancia de la fotosintesis
    document.add_heading('Importancia de la Fotosintesis', level=1)
    document.add_paragraph('La fotosintesis es esencial para la vida en la Tierra porque:')
    document.add_paragraph(' - Produce oxigeno, esencial para la respiracion de la mayoria de los organismos.')
    document.add_paragraph(' - Es la base de la cadena alimentaria, proporcionando energia a la mayoria de los ecosistemas.')
    document.add_paragraph(' - Ayuda a regular el ciclo del carbono en la atmosfera.')


    # Guardar el documento
    document.save('trabajo_fotosintesis.docx')

    # Abrir el documento (opcional - requiere que Word este instalado)
    try:
        os.startfile('trabajo_fotosintesis.docx')  #Para Windows
        #Para MacOS o Linux, deberas usar un comando diferente segun tu sistema operativo.
    except:
        print('El archivo se creo correctamente, pero no se pudo abrir automaticamente. Por favor, abrelo manualmente.')



#Ejecutar la funcion
crear_trabajo_fotosintesis()

